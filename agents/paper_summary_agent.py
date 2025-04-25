import fitz  # For reading PDFs
import streamlit as st
import docx
from datetime import datetime
from langchain.prompts import PromptTemplate


def run(llm):
    st.title("📄 Paper Summary Agent")

    uploaded_files = st.file_uploader(
        "Upload one or more scientific papers (PDF)",
        type=["pdf"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        summaries = []

        for uploaded_file in uploaded_files:
            # Read the PDF
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                text = "\n".join(page.get_text() for page in doc)

            # Prepare the structured summarization prompt
            prompt = PromptTemplate(
                input_variables=["content"],
                template="""You are an expert at reading and summarizing scientific papers.
Given the following full text of a paper, extract the following:
- Paper Title
- Authors
- Year
- Summary
- Methodology
- Results
- Conclusion and Future Work

Respond STRICTLY in this format (no extra explanations):

Title: <text>
Authors: <text>
Year: <text>
Summary: <text>
Methodology: <text>
Results: <text>
Conclusion and Future Work: <text>

Paper Text:
{content}
""",
            )

            # Apply prompt to the text (limit if necessary)
            formatted_prompt = prompt.format(
                content=text[:12000]
            )  # limit to avoid token overflows
            response = llm.predict(formatted_prompt)
            summaries.append(response)

        # === Create a Word Document ===
        doc = docx.Document()
        doc.add_heading("Paper Summaries", level=0)

        for idx, summary in enumerate(summaries, 1):
            doc.add_heading(f"Paper {idx}", level=1)
            for line in summary.split("\n"):
                if line.strip():
                    if any(
                        section in line for section in ["Title:", "Authors:", "Year:"]
                    ):
                        doc.add_heading(line.strip(), level=2)
                    else:
                        doc.add_paragraph(line.strip())

        # Save Word document with timestamp to avoid overwriting
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"paper_summaries_{timestamp}.docx"
        output_path = f"data/outputs/{output_filename}"
        doc.save(output_path)

        # === Offer Download Button ===
        with open(output_path, "rb") as f:
            st.download_button(
                label="📄 Download Paper Summaries",
                data=f,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
